# -*- coding: utf-8 -*-
"""
Build final thesis DOCX with all images embedded.
- Replaces IMAGE PLACEHOLDER paragraphs with real screenshots
- Inserts diagram images before their figure captions
- Adds section 4.5 mobile screenshots
- Adds Appendix D: Wireframes
- Inserts dynamic TOC field (auto-updates when opened in Word)
"""
import pathlib, copy, shutil
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import lxml.etree as etree

BASE   = pathlib.Path(__file__).parent
DDDD   = BASE / 'dddddd'
WEB    = DDDD / 'web_screenshots'
MOB    = DDDD / 'mobile'
WIRE   = DDDD / 'wireframes'
DIAG   = DDDD / 'diagrams'
SRC    = DDDD / 'Graduation Project Shamel.docx'
OUT    = DDDD / 'Graduation_Project_Shamel_FINAL.docx'

IMG_W  = Inches(5.8)   # image width that fits A4 with margins
CAP_W  = Inches(5.8)

# ── helpers ──────────────────────────────────────────────────────────────────

def add_page_break(doc):
    p = doc.add_paragraph()
    run = p.add_run()
    run.add_break(docx_break_type_page())
    return p

def docx_break_type_page():
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    br = OxmlElement('w:br')
    br.set(qn('w:type'), 'page')
    return br

def insert_page_break_before(paragraph):
    """Add page break at start of paragraph via pPr > pageBreakBefore."""
    pPr = paragraph._p.get_or_add_pPr()
    pb = OxmlElement('w:pageBreakBefore')
    pb.set(qn('w:val'), '1')
    pPr.append(pb)

def add_image_paragraph(doc, img_path, width=IMG_W):
    """Add a centered paragraph containing one image. Returns the paragraph."""
    if not pathlib.Path(img_path).exists():
        p = doc.add_paragraph(f'[IMAGE NOT FOUND: {img_path}]')
        return p
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(img_path), width=width)
    return p

def add_caption(doc, text, style_name='Caption'):
    """Add a figure caption paragraph."""
    try:
        p = doc.add_paragraph(style=style_name)
    except KeyError:
        p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = False
    run.font.size = Pt(10)
    return p

def set_arabic_caption(p, text):
    p.clear()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run(text)
    run.font.size = Pt(10)
    run.font.name = 'Simplified Arabic'

def para_index(doc, para):
    return list(doc.paragraphs).index(para)

def insert_paragraph_before(ref_para, doc):
    """Insert a new empty paragraph before ref_para. Returns new paragraph."""
    new_p = OxmlElement('w:p')
    ref_para._p.addprevious(new_p)
    from docx.text.paragraph import Paragraph
    return Paragraph(new_p, ref_para._p.getparent())

def insert_image_before_para(ref_para, img_path, doc, width=IMG_W):
    """Insert page-break + image paragraph immediately before ref_para."""
    if not pathlib.Path(img_path).exists():
        return
    # Create real paragraph at end of doc (so it has doc part), then move it
    img_para = doc.add_paragraph()
    img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = img_para.add_run()
    run.add_picture(str(img_path), width=width)
    # Move the XML element to just before ref_para
    ref_para._p.addprevious(img_para._p)

    # Page break paragraph (raw XML, no picture needed)
    pb_p = OxmlElement('w:p')
    pb_r = OxmlElement('w:r')
    pb_br = OxmlElement('w:br')
    pb_br.set(qn('w:type'), 'page')
    pb_r.append(pb_br)
    pb_p.append(pb_r)
    img_para._p.addprevious(pb_p)
    return img_para

def add_toc_field(doc):
    """Replace TOC placeholder with a Word TOC field that auto-updates."""
    # Find the TABLE OF CONTENTS paragraphs
    for i, para in enumerate(doc.paragraphs):
        if para.text.strip() == 'TABLE OF CONTENTS':
            # Clear everything after TOC heading until next section
            # Insert TOC field code
            toc_para = doc.add_paragraph()
            toc_para._p.getparent().remove(toc_para._p)

            # Build the field XML
            fld_xml = (
                '<w:p xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
                '<w:pPr><w:jc w:val="center"/></w:pPr>'
                '<w:fldSimple w:instr=" TOC \\o &quot;1-3&quot; \\h \\z \\u ">'
                '<w:r><w:t>Right-click and select "Update Field" to generate Table of Contents</w:t></w:r>'
                '</w:fldSimple>'
                '</w:p>'
            )
            fld_elem = etree.fromstring(fld_xml)

            # Insert after the TOC heading paragraph
            next_p = para._p.getnext()
            if next_p is not None:
                next_p.addprevious(fld_elem)
            else:
                para._p.getparent().append(fld_elem)
            break

# ── image mappings ────────────────────────────────────────────────────────────

# Placeholder text → image file
PLACEHOLDER_MAP = {
    'login.html':                    WEB / 'web_01_login.png',
    'face_login.html':               WEB / 'web_02_face_login.png',
    'enroll_face.html':              WEB / 'web_14_enroll_face.png',
    'exam_gate_verify.html':         WEB / 'web_32_exam_gate.png',
    'onboarding_wizard.html':        WEB / 'web_13_onboarding.png',
    'coordinator_dashboard.html':    WEB / 'web_15_coord_dash.png',
    'coordinator_students.html':     WEB / 'web_16_coord_students.png',
    'coordinator_faculty.html':      WEB / 'web_17_coord_faculty.png',
    'coordinator_assignments.html':  WEB / 'web_18_coord_assign.png',
}

# Figure caption text → image file (for figures WITHOUT placeholders)
CAPTION_TO_IMAGE = {
    '4.4.3':  WEB / 'web_01_login.png',         # SSO Portal - use login
    '4.4.5':  WEB / 'web_29_gate_dash.png',      # scan → gate as proxy
    '4.4.6':  WEB / 'web_29_gate_dash.png',      # gate
    '4.4.8':  WEB / 'web_03_admin_dash.png',     # admin panel (note: Figure 4.8 not 4.4.8)
    '4.8':    WEB / 'web_03_admin_dash.png',
    '4.4.9':  WEB / 'web_04_gate_logs.png',
    '4.4.10': WEB / 'web_11_notifications.png',
    '4.4.11': WEB / 'web_10_audit.png',
    '4.4.17': WEB / 'web_06_students.png',
    '4.4.18': WEB / 'web_31_teacher_detail.png',
    '4.1':    None,   # ROC chart - skip
    '2.1':    None,   # theoretical schematic - skip
    '3.1':    DIAG / 'fig3_1_architecture.png',
    '3.2':    DIAG / 'fig3_2_erd.png',
    '3.3':    DIAG / 'fig3_3_usecase.png',
    '3.4':    DIAG / 'fig3_4_sequence.png',
    '3.5':    None,   # activity flow - not generated
}

# Mobile screenshots for section 4.5
MOBILE_PAGES = [
    (MOB / 'mob_01_launch.png',          'شكل 4.5.1: شاشة إطلاق التطبيق — واجهة البداية مع شعار شامل'),
    (MOB / 'mob_02_login_filled.png',    'شكل 4.5.2: شاشة تسجيل الدخول — إدخال بيانات الاعتماد'),
    (MOB / 'mob_03_student_dashboard.png','شكل 4.5.3: لوحة الطالب الرئيسية — ملخص الحضور والمقررات'),
    (MOB / 'mob_04_schedule.png',        'شكل 4.5.4: الجدول الدراسي للطالب — عرض الجدول الأسبوعي'),
    (MOB / 'mob_06_drawer_open.png',     'شكل 4.5.5: القائمة الجانبية — التنقل بين أقسام التطبيق'),
    (MOB / 'mob_10_attendance_screen.png','شكل 4.5.6: شاشة سجل الحضور — تفاصيل حضور المادة الدراسية'),
    (MOB / 'admin_home.png',             'شكل 4.5.7: لوحة المدير في التطبيق — إحصاءات الجامعة الشاملة'),
    (MOB / 'admin_tab_reports.png',      'شكل 4.5.8: تبويب التقارير للمدير — تقارير الأداء والحضور'),
    (MOB / 'teacher_home.png',           'شكل 4.5.9: لوحة الأستاذ — جلسات المحاضرات والحضور'),
    (MOB / 'coord_home.png',             'شكل 4.5.10: لوحة منسق الكلية — مؤشرات الأداء الأكاديمي'),
    (MOB / 'gate_home.png',              'شكل 4.5.11: لوحة حارس البوابة — مراقبة دخول الحرم'),
]

# Wireframes for appendix
WF_WEB_PAGES = sorted(WIRE.glob('wf_web_*.png'))
WF_MOB_PAGES = sorted(WIRE.glob('wf_mob_*.png'))
WF_LEGACY    = sorted(WIRE.glob('wf0*.png'))

# Wireframe descriptions (by filename prefix)
def wf_desc(fname):
    desc_map = {
        'wf_web_01': 'إطار سلكي — صفحة تسجيل الدخول الرئيسية',
        'wf_web_02': 'إطار سلكي — صفحة إعادة تعيين كلمة المرور',
        'wf_web_03': 'إطار سلكي — لوحة تحكم المدير',
        'wf_web_04': 'إطار سلكي — سجلات بوابة الدخول',
        'wf_web_05': 'إطار سلكي — إدارة أعضاء هيئة التدريس',
        'wf_web_06': 'إطار سلكي — إدارة الطلاب',
        'wf_web_07': 'إطار سلكي — إدارة المقررات',
        'wf_web_08': 'إطار سلكي — الجدول الدراسي',
        'wf_web_09': 'إطار سلكي — التقويم الجدولي',
        'wf_web_10': 'إطار سلكي — إدارة القاعات',
        'wf_web_11': 'إطار سلكي — التقارير والإحصاءات',
        'wf_web_12': 'إطار سلكي — التحليلات المتقدمة',
        'wf_web_13': 'إطار سلكي — البحث في النظام',
        'wf_web_14': 'إطار سلكي — إعدادات النظام',
        'wf_web_15': 'إطار سلكي — مركز الإشعارات',
        'wf_web_16': 'إطار سلكي — سجل التدقيق',
        'wf_web_17': 'إطار سلكي — لوحة أستاذ المادة',
        'wf_web_18': 'إطار سلكي — الجلسات الدراسية',
        'wf_web_19': 'إطار سلكي — تسجيل الحضور',
        'wf_web_20': 'إطار سلكي — الخط الزمني للأستاذ',
        'wf_web_21': 'إطار سلكي — سجلات الحضور',
        'wf_web_22': 'إطار سلكي — جدول الأستاذ',
        'wf_web_23': 'إطار سلكي — لوحة الطالب',
        'wf_web_24': 'إطار سلكي — مقررات الطالب',
        'wf_web_25': 'إطار سلكي — جدول الطالب',
        'wf_web_26': 'إطار سلكي — بوابة الأعذار',
        'wf_web_27': 'إطار سلكي — درجات الطالب',
        'wf_web_28': 'إطار سلكي — إشعارات الطالب',
        'wf_web_29': 'إطار سلكي — ملف الطالب الشخصي',
        'wf_web_30': 'إطار سلكي — لوحة منسق الكلية',
        'wf_web_31': 'إطار سلكي — طلاب الكلية',
        'wf_web_32': 'إطار سلكي — الكادر الأكاديمي',
        'wf_web_33': 'إطار سلكي — الدرجات والتقييم',
        'wf_web_34': 'إطار سلكي — تسجيل المستخدمين',
        'wf_web_35': 'إطار سلكي — لوحة حارس البوابة',
        'wf_web_36': 'إطار سلكي — مسح الوجه في البوابة',
        'wf_web_37': 'إطار سلكي — سجلات البوابة',
        'wf_web_38': 'إطار سلكي — تذاكر الدعم الفني',
        'wf_web_39': 'إطار سلكي — تسجيل بصمة الوجه',
        'wf_mob_01': 'إطار سلكي موبايل — شاشة تسجيل الدخول',
        'wf_mob_02': 'إطار سلكي موبايل — لوحة المدير',
        'wf_mob_03': 'إطار سلكي موبايل — جدول المدير',
        'wf_mob_04': 'إطار سلكي موبايل — تقارير المدير',
        'wf_mob_05': 'إطار سلكي موبايل — ملف المدير',
        'wf_mob_06': 'إطار سلكي موبايل — تسجيل الطلاب',
        'wf_mob_07': 'إطار سلكي موبايل — لوحة الأستاذ',
        'wf_mob_08': 'إطار سلكي موبايل — جدول الأستاذ',
        'wf_mob_09': 'إطار سلكي موبايل — تقارير الأستاذ',
        'wf_mob_10': 'إطار سلكي موبايل — ملف الأستاذ',
        'wf_mob_11': 'إطار سلكي موبايل — لوحة الطالب',
        'wf_mob_12': 'إطار سلكي موبايل — جدول الطالب',
        'wf_mob_13': 'إطار سلكي موبايل — تقارير الطالب',
        'wf_mob_14': 'إطار سلكي موبايل — ملف الطالب',
        'wf_mob_15': 'إطار سلكي موبايل — لوحة البوابة',
    }
    stem = fname.stem
    return desc_map.get(stem, f'إطار سلكي — {stem}')

# ── main build ────────────────────────────────────────────────────────────────

def build():
    shutil.copy(str(SRC), str(OUT))
    doc = Document(str(OUT))

    paragraphs = doc.paragraphs  # live list

    # ── Pass 1: Replace IMAGE PLACEHOLDER paragraphs ─────────────────────────
    to_replace = []  # list of (para_index, img_path)
    for i, para in enumerate(paragraphs):
        txt = para.text.strip()
        if '[IMAGE INTERFACE PLACEHOLDER:' in txt:
            matched_img = None
            for key, img in PLACEHOLDER_MAP.items():
                if key in txt:
                    matched_img = img
                    break
            to_replace.append((para, matched_img))

    for (para, img_path) in to_replace:
        # Clear the placeholder text
        for run in para.runs:
            run.text = ''
        para.text = ''
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if img_path and pathlib.Path(img_path).exists():
            run = para.add_run()
            run.add_picture(str(img_path), width=IMG_W)
        else:
            para.add_run(f'[Image unavailable]')
        # Page break before this image
        insert_page_break_before(para)

    print(f'Replaced {len(to_replace)} image placeholders')

    # ── Pass 2: Insert images before Figure captions (no placeholder) ─────────
    inserted = 0
    for para in list(doc.paragraphs):
        txt = para.text.strip()
        if not txt.startswith('Figure'):
            continue
        # Extract figure number
        fig_img = None
        for key, img in CAPTION_TO_IMAGE.items():
            # Match "Figure 3.1:" or "Figure 4.4.9:"
            if f'Figure {key}:' in txt or f'Figure {key} ' in txt:
                fig_img = img
                break
        if fig_img is None:
            continue
        if fig_img is not None and pathlib.Path(fig_img).exists():
            insert_image_before_para(para, fig_img, doc, width=IMG_W)
            inserted += 1

    print(f'Inserted {inserted} images before figure captions')

    # ── Pass 3: Add mobile screenshots into section 4.5 ──────────────────────
    # Find the "4.5 Comprehensive System Interfaces" paragraph
    sec45_para = None
    for para in doc.paragraphs:
        if '4.5 Comprehensive System Interfaces' in para.text:
            sec45_para = para
            break

    if sec45_para:
        # Find the paragraph AFTER 4.5 text body (before 4.6)
        # Use element-based search (robust after prior insertions)
        paras = list(doc.paragraphs)
        sec45_elem = sec45_para._p
        try:
            idx45 = next(i for i, p in enumerate(paras) if p._p is sec45_elem)
        except StopIteration:
            idx45 = 0
        # Find where 4.6 starts
        insert_before = None
        for p in paras[idx45+1:]:
            if p.text.strip().startswith('4.6'):
                insert_before = p
                break

        if insert_before is None:
            # Just append to end of doc
            target = doc
            for (img_path, caption_ar) in MOBILE_PAGES:
                if img_path.exists():
                    # page break
                    pb = doc.add_paragraph()
                    pb.add_run().add_break(docx_break_type_page())
                    add_image_paragraph(doc, img_path)
                    cap = doc.add_paragraph()
                    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    r = cap.add_run(caption_ar)
                    r.font.size = Pt(10)
                    r.font.name = 'Simplified Arabic'
        else:
            # Insert before "4.6" paragraph (reverse order so first ends up first)
            for (img_path, caption_ar) in reversed(MOBILE_PAGES):
                if not img_path.exists():
                    continue
                # Caption — real paragraph moved into position
                cap_para = doc.add_paragraph()
                cap_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                r = cap_para.add_run(caption_ar)
                r.font.size = Pt(10)
                r.font.name = 'Simplified Arabic'
                insert_before._p.addprevious(cap_para._p)
                # Image
                img_para = doc.add_paragraph()
                img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                img_para.add_run().add_picture(str(img_path), width=Inches(3.0))
                cap_para._p.addprevious(img_para._p)
                # Page break
                pb_elem = OxmlElement('w:p')
                pb_r = OxmlElement('w:r')
                pb_br = OxmlElement('w:br')
                pb_br.set(qn('w:type'), 'page')
                pb_r.append(pb_br)
                pb_elem.append(pb_r)
                img_para._p.addprevious(pb_elem)

        print(f'Added {len(MOBILE_PAGES)} mobile screenshots to section 4.5')

    # ── Pass 4: Add Wireframes Appendix ──────────────────────────────────────
    # Find APPENDICES section at end
    appendix_para = None
    for para in doc.paragraphs:
        if 'APPENDICES' in para.text and 'APPENDIX A' in para.text:
            appendix_para = para
            break
    if not appendix_para:
        for para in doc.paragraphs:
            if para.text.strip() == 'APPENDICES':
                appendix_para = para
                break

    # Add Appendix D: Wireframes after existing appendices
    # Append at end of document
    doc.add_page_break()
    h = doc.add_paragraph()
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = h.add_run('APPENDIX D: SYSTEM WIREFRAMES AND UI PROTOTYPES')
    r.bold = True
    r.font.size = Pt(12)
    r.font.name = 'Times New Roman'

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rs = sub.add_run('الإطارات السلكية وتصميمات واجهات النظام')
    rs.bold = True
    rs.font.size = Pt(14)
    rs.font.name = 'Simplified Arabic'

    doc.add_paragraph()

    intro = doc.add_paragraph()
    intro.alignment = WD_ALIGN_PARAGRAPH.LEFT
    ri = intro.add_run(
        'This appendix presents the complete collection of wireframe prototypes '
        'developed during the design phase of shamel. Wireframes cover all five user roles '
        '(Admin, Coordinator, Teacher, Student, Gate Operator) across both the web platform '
        'and the mobile application. Each wireframe represents the structural layout and '
        'functional flow of a single interface screen.'
    )
    ri.font.size = Pt(12)
    ri.font.name = 'Times New Roman'

    # D.1 Web Wireframes
    doc.add_page_break()
    sh = doc.add_paragraph()
    sh.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rsh = sh.add_run('D.1  Web Application Wireframes')
    rsh.bold = True
    rsh.font.size = Pt(12)

    wf_fig_num = 1
    for wf_file in WF_WEB_PAGES:
        if not wf_file.exists():
            continue
        doc.add_page_break()
        ip = doc.add_paragraph()
        ip.alignment = WD_ALIGN_PARAGRAPH.CENTER
        ip.add_run().add_picture(str(wf_file), width=IMG_W)
        cp = doc.add_paragraph()
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        rc = cp.add_run(f'Figure D.{wf_fig_num}: {wf_desc(wf_file)}')
        rc.font.size = Pt(10)
        wf_fig_num += 1

    # D.2 Mobile Wireframes
    doc.add_page_break()
    sh2 = doc.add_paragraph()
    sh2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rsh2 = sh2.add_run('D.2  Mobile Application Wireframes')
    rsh2.bold = True
    rsh2.font.size = Pt(12)

    for wf_file in WF_MOB_PAGES:
        if not wf_file.exists():
            continue
        doc.add_page_break()
        ip = doc.add_paragraph()
        ip.alignment = WD_ALIGN_PARAGRAPH.CENTER
        ip.add_run().add_picture(str(wf_file), width=Inches(3.0))
        cp = doc.add_paragraph()
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        rc = cp.add_run(f'Figure D.{wf_fig_num}: {wf_desc(wf_file)}')
        rc.font.size = Pt(10)
        wf_fig_num += 1

    print(f'Added wireframes appendix ({wf_fig_num - 1} wireframes)')

    # ── Pass 5: Fix TOC to use Word field ─────────────────────────────────────
    # Find and update the two TOC placeholder paragraphs
    toc_paras = []
    for para in doc.paragraphs:
        t = para.text.strip()
        if t in ('TABLE OF CONTENTS', 'TABLE OF CONTENTS (CONTINUED)'):
            toc_paras.append(para)

    # For the first TOC paragraph, add a TOC field instruction after it
    if toc_paras:
        first_toc = toc_paras[0]
        # Insert a new paragraph with TOC field after the heading
        toc_field_xml = (
            '<w:p xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
            '<w:pPr><w:jc w:val="left"/></w:pPr>'
            '<w:fldSimple w:instr=" TOC \\o &quot;1-3&quot; \\h \\z \\u ">'
            '<w:r><w:rPr><w:lang w:val="en-US"/></w:rPr>'
            '<w:t>[ Open in Microsoft Word and press Ctrl+A then F9 to update Table of Contents ]</w:t>'
            '</w:r>'
            '</w:fldSimple>'
            '</w:p>'
        )
        toc_field_elem = etree.fromstring(toc_field_xml)
        first_toc._p.addnext(toc_field_elem)
        print('TOC field inserted')

    # ── Save ──────────────────────────────────────────────────────────────────
    doc.save(str(OUT))
    print(f'\nSaved: {OUT}')
    size_mb = OUT.stat().st_size / 1024 / 1024
    print(f'File size: {size_mb:.1f} MB')


if __name__ == '__main__':
    build()
